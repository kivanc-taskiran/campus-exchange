"""
Generates ACM474_Term_Homework.docx for Kivanc Mete Taskiran.
Run:  python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_row_bg(row, hex_color):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return p

def add_body(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p

def add_code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x21)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F3F4F6")
        pPr.append(shd)

def add_output_block(doc, output):
    for line in output.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xD1, 0xFA, 0xE5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F2937")
        pPr.append(shd)

def add_table(doc, headers, rows, header_color="1A56DB", alt_color="EFF6FF"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    set_row_bg(hdr_row, header_color)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx+1]
        if r_idx % 2 == 1:
            set_row_bg(row, alt_color)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ── Document Setup ─────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# Cover
doc.add_paragraph()
title = doc.add_heading("ACM 474 \u2013 Term Homework", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
sub = doc.add_paragraph("K\u0131van\u00e7 Mete Ta\u015fk\u0131ran")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(14)
sub2 = doc.add_paragraph("Date: June 2026")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── PART 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, "PART 1 (3 Points) \u2014 Caesar Cipher & Vigenere Cipher Programs", 1)
add_heading(doc, "Parameters", 2)
add_table(doc, ["Parameter","Value"], [
    ["Plaintext",             "MYNAMEISKIVANCMETETASKIRAN"],
    ["Name (Latin alphabet)", "KIVANC METE TASKIRAN"],
    ["Caesar Shift Key",      "18  (6+4+8 = total chars in KIVANC+METE+TASKIRAN)"],
    ["Vigenere Key",          "SECURITY"],
])

# 1A Caesar
add_heading(doc, "Program 1A \u2014 Caesar (Shift) Cipher", 2)
add_code_block(doc, """\
# ACM 474 - Term Homework | Part 1: Caesar (Shift) Cipher
# Student : Kivanc Mete Taskiran
# Plaintext: MYNAMEISKIVANCMETETASKIRAN
# Shift    : 18  (total chars in KIVANC + METE + TASKIRAN)

def encrypt(text, s):
    result = ""
    for i in range(len(text)):
        char = text[i]
        # Encrypt uppercase characters
        if char.isupper():
            result += chr((ord(char) + s - 65) % 26 + 65)
        # Encrypt lowercase characters
        else:
            result += chr((ord(char) + s - 97) % 26 + 97)
    return result


def decrypt(text, s):
    result = ""
    for i in range(len(text)):
        char = text[i]
        # Decrypt uppercase characters
        if char.isupper():
            result += chr((ord(char) - s - 65) % 26 + 65)
        # Decrypt lowercase characters
        else:
            result += chr((ord(char) - s - 97) % 26 + 97)
    return result


# Driver code
text = "MYNAMEISKIVANCMETETASKIRAN"
s = 18  # len("KIVANC") + len("METE") + len("TASKIRAN") = 6 + 4 + 8 = 18

print("Plaintext:  " + text)
print("Shift:      " + str(s))
print("Ciphertext: " + encrypt(text, s))
print("Decrypted:  " + decrypt(encrypt(text, s), s))""")
doc.add_paragraph()
add_body(doc, "Output:", bold=True)
add_output_block(doc, """\
Plaintext:  MYNAMEISKIVANCMETETASKIRAN
Shift:      18
Ciphertext: EQFSEWAKCANSFUEWLWLSKCAJSF
Decrypted:  MYNAMEISKIVANCMETETASKIRAN""")
doc.add_paragraph()

# 1B Vigenere
add_heading(doc, "Program 1B \u2014 Vigenere Cipher", 2)
add_code_block(doc, """\
# ACM 474 - Term Homework | Part 1: Vigenere Cipher
# Student : Kivanc Mete Taskiran
# Plaintext: MYNAMEISKIVANCMETETASKIRAN
# Key      : SECURITY

# This function generates the key in a cyclic manner until
# its length isn't equal to the length of original text
def generateKey(string, key):
    key = list(key)
    if len(string) == len(key):
        return(key)
    else:
        for i in range(len(string) -
                       len(key)):
            key.append(key[i % len(key)])
    return("" . join(key))

# This function returns the encrypted text generated
# with the help of the key
def cipherText(string, key):
    cipher_text = []
    for i in range(len(string)):
        x = (ord(string[i]) +
             ord(key[i])) % 26
        x += ord('A')
        cipher_text.append(chr(x))
    return("" . join(cipher_text))

# This function decrypts the encrypted text and returns
# the original text
def originalText(cipher_text, key):
    orig_text = []
    for i in range(len(cipher_text)):
        x = (ord(cipher_text[i]) -
             ord(key[i]) + 26) % 26
        x += ord('A')
        orig_text.append(chr(x))
    return("" . join(orig_text))

# Driver code
if __name__ == "__main__":
    string  = "MYNAMEISKIVANCMETETASKIRAN"
    keyword = "SECURITY"
    key = generateKey(string, keyword)
    cipher_text = cipherText(string, key)
    print("Plaintext :              " + string)
    print("Key :                    " + keyword)
    print("Ciphertext :             " + cipher_text)
    print("Original/Decrypted Text: " + originalText(cipher_text, key))""")
doc.add_paragraph()
add_body(doc, "Output:", bold=True)
add_output_block(doc, """\
Plaintext :              MYNAMEISKIVANCMETETASKIRAN
Key :                    SECURITY
Ciphertext :             ECPUDMBQCMXUEKFCLIVUJSBPSR
Original/Decrypted Text: MYNAMEISKIVANCMETETASKIRAN""")

doc.add_page_break()

# ── PART 2 ─────────────────────────────────────────────────────────────────
add_heading(doc, "PART 2 (3 Points) \u2014 Combined Caesar + Vigenere Cipher", 1)
add_code_block(doc, """\
# ACM 474 - Term Homework | Part 2: Combined Caesar + Vigenere
# Student : Kivanc Mete Taskiran
# Plaintext   : MYNAMEISKIVANCMETETASKIRAN
# Caesar Shift: 18  (len("KIVANC")+len("METE")+len("TASKIRAN"))
# Vigenere Key: SECURITY


# ── Caesar Cipher Functions ────────────────────────────────

def encrypt(text, s):
    result = ""
    for i in range(len(text)):
        char = text[i]
        if char.isupper():
            result += chr((ord(char) + s - 65) % 26 + 65)
        else:
            result += chr((ord(char) + s - 97) % 26 + 97)
    return result


def decrypt(text, s):
    result = ""
    for i in range(len(text)):
        char = text[i]
        if char.isupper():
            result += chr((ord(char) - s - 65) % 26 + 65)
        else:
            result += chr((ord(char) - s - 97) % 26 + 97)
    return result


# ── Vigenere Cipher Functions ──────────────────────────────

# This function generates the key in a cyclic manner until
# its length isn't equal to the length of original text
def generateKey(string, key):
    key = list(key)
    if len(string) == len(key):
        return(key)
    else:
        for i in range(len(string) -
                       len(key)):
            key.append(key[i % len(key)])
    return("" . join(key))

# This function returns the encrypted text generated
# with the help of the key
def cipherText(string, key):
    cipher_text = []
    for i in range(len(string)):
        x = (ord(string[i]) +
             ord(key[i])) % 26
        x += ord('A')
        cipher_text.append(chr(x))
    return("" . join(cipher_text))

# This function decrypts the encrypted text and returns
# the original text
def originalText(cipher_text, key):
    orig_text = []
    for i in range(len(cipher_text)):
        x = (ord(cipher_text[i]) -
             ord(key[i]) + 26) % 26
        x += ord('A')
        orig_text.append(chr(x))
    return("" . join(orig_text))


# ── Driver Code ────────────────────────────────────────────

if __name__ == "__main__":
    plaintext = "MYNAMEISKIVANCMETETASKIRAN"
    s         = 18           # Caesar shift key
    keyword   = "SECURITY"   # Vigenere key

    print("Plaintext: " + plaintext)

    # STEP 1: Apply Caesar encryption on the plaintext
    ciphertext1 = encrypt(plaintext, s)
    print("After Caesar Encryption (ciphertext1):   " + ciphertext1)

    # STEP 2: Apply Vigenere encryption on the Caesar ciphertext
    key = generateKey(ciphertext1, keyword)
    ciphertext2 = cipherText(ciphertext1, key)
    print("After Vigenere Encryption (ciphertext2): " + ciphertext2)

    # STEP 3: Decrypt - Vigenere first, then Caesar
    vigenere_key        = generateKey(ciphertext2, keyword)
    vigenere_decrypted  = originalText(ciphertext2, vigenere_key)
    decrypted_plaintext = decrypt(vigenere_decrypted, s)
    print("Original/Decrypted Text: " + decrypted_plaintext)""")
doc.add_paragraph()
add_body(doc, "Output:", bold=True)
add_output_block(doc, """\
Plaintext: MYNAMEISKIVANCMETETASKIRAN
After Caesar Encryption (ciphertext1):   EQFSEWAKCANSFUEWLWLSKCAJSF
After Vigenere Encryption (ciphertext2): WUHMVETIUEPMWCXUDANMBKTHKJ
Original/Decrypted Text: MYNAMEISKIVANCMETETASKIRAN""")

doc.add_page_break()

# ── PART 3 ─────────────────────────────────────────────────────────────────
add_heading(doc, "PART 3 (4 Points) \u2014 Client-Server Network Design (1,000 Employees)", 1)
add_body(doc,
    "The following describes a professional client-server network infrastructure suitable for a "
    "medium-to-large IT company with 1,000 employees. The design prioritizes performance, security, "
    "scalability, and redundancy \u2014 all critical requirements for a company in the IT sector."
)
doc.add_paragraph()

add_heading(doc, "1. How Many Client Computers Would Be Needed?", 2)
add_body(doc,
    "A company with 1,000 employees would require approximately 1,100\u20131,130 client computers:\n"
    "\u2022 1,000 primary workstations \u2014 one per employee (desktops or laptops depending on role)\n"
    "\u2022 50\u2013100 spare/reserve machines for replacements, temporary staff, and meeting rooms\n"
    "\u2022 20\u201330 thin clients placed in shared areas (reception, conference rooms, lounges)\n\n"
    "For an IT company, many employees would also use personal laptops and mobile devices (BYOD policy), "
    "registered on the network under a Mobile Device Management (MDM) system."
)

add_heading(doc, "2. How Many Server Computers Would Be Needed?", 2)
add_body(doc, "A 1,000-person IT company would require a minimum of 25\u201330 dedicated physical servers:")
add_table(doc, ["Server Type","Count","Purpose"], [
    ["Application / Web Servers",           "3\u20134","Internal web apps, intranet, customer-facing services"],
    ["Database Servers",                    "3\u20134","Relational and NoSQL databases"],
    ["File / Storage Servers",              "2\u20133","Centralized file storage, document management"],
    ["Mail Servers",                        "2",       "Corporate email (Exchange, Postfix)"],
    ["Authentication / Directory Servers",  "2",       "Active Directory / LDAP for user logins"],
    ["Backup Servers",                      "2",       "Automated data backup and disaster recovery"],
    ["Virtualization / Hypervisor Servers", "3\u20134","Host VMs for dev, test, and staging environments"],
    ["Proxy / Load Balancer Servers",       "2",       "Traffic distribution and internet access management"],
    ["Monitoring / Logging Servers",        "1\u20132","System health, performance, and log collection"],
    ["Communication Servers",              "1\u20132","VoIP, video conferencing, internal chat"],
    ["DNS / DHCP Servers",                  "2",       "Dynamic IP addressing and domain name resolution"],
    ["VPN Servers",                         "2",       "Secure remote access for off-site employees"],
])
add_body(doc, "A hybrid cloud model (Microsoft Azure, AWS) would supplement physical servers for overflow capacity and disaster recovery.")

add_heading(doc, "3. How Many Shared Peripheral Devices Would Be Needed?", 2)
add_table(doc, ["Device Type","Count","Rationale"], [
    ["Network Printers (Multifunction)", "30\u201340",   "1 per ~25 employees, shared by department"],
    ["Scanners",                         "10\u201315",   "1 per ~70 employees for document digitization"],
    ["Photocopiers",                     "10\u201315",   "Shared per floor or department"],
    ["NAS Devices",                      "5\u201310",    "Departmental shared file storage"],
    ["Projectors / Smart Displays",      "20\u201330",   "1\u20132 per conference/meeting room"],
    ["VoIP Phone Handsets",              "200\u2013300", "For departments still using desk phones"],
    ["Shared UPS Units",                 "10\u201315",   "Server rooms and critical network equipment"],
])
add_body(doc, "Total peripheral devices: approximately 300\u2013400 units.")

add_heading(doc, "4. Network Topology and Layering Model", 2)
add_body(doc,
    "Topology: A Hierarchical Star (Hybrid) Topology would be used \u2014 the industry standard for "
    "enterprise networks of this scale, consisting of three layers:\n"
    "\u2022 Core Layer: Two high-performance redundant core switches (10 GbE or higher)\n"
    "\u2022 Distribution Layer: Multiple distribution switches per floor aggregating access-layer traffic\n"
    "\u2022 Access Layer: Switches at each floor connecting client machines and peripheral devices\n\n"
    "Layering Model: The 5-Layer TCP/IP Hybrid Model (Physical, Data Link, Network, Transport, Application) "
    "would be used. It is preferred over the 7-layer OSI model (theoretical reference only) and the 4-layer "
    "TCP/IP model because it maps most accurately to real-world network implementations."
)

add_heading(doc, "5. Wired and Wireless Media", 2)
add_body(doc, "Wired Media:", bold=True)
add_table(doc, ["Media Type","Usage"], [
    ["Cat6A UTP Cables",          "Primary client-to-switch connections (supports 10 GbE up to 100m)"],
    ["Cat6 UTP Cables",           "General office workstation connections"],
    ["Fiber-Optic (Single-mode)", "Backbone links between core/distribution switches and inter-building"],
    ["Fiber-Optic (Multi-mode)",  "Short-distance high-speed links within the server room"],
])
add_body(doc, "Wireless Media:", bold=True)
add_table(doc, ["Standard","Usage"], [
    ["IEEE 802.11ax (Wi-Fi 6/6E)","Primary wireless for laptops and mobile devices; high capacity, low latency"],
    ["IEEE 802.11ac (Wi-Fi 5)",   "Legacy device support and lower-traffic areas"],
    ["IEEE 802.15.1 (Bluetooth)", "Short-range device pairing (keyboards, mice, headsets)"],
])
add_body(doc,
    "WAPs would be deployed across all floors with a centralized WLAN controller. "
    "WPA3-Enterprise (IEEE 802.1X + RADIUS) would be used for wireless security."
)

add_heading(doc, "6. Network Operating System (NOS)", 2)
add_table(doc, ["OS","Role"], [
    ["Windows Server 2022",          "Active Directory, file sharing, Group Policy, DNS/DHCP, Microsoft 365 integration"],
    ["Linux (Ubuntu Server / RHEL)", "Web servers, databases, dev/test environments \u2014 stability and lower cost"],
    ["VMware ESXi / Hyper-V",        "Hypervisor OS for managing virtual machines"],
])

add_heading(doc, "7. Firewalls, Security Servers, and Anti-Virus Software", 2)
add_table(doc, ["Firewall Type","Placement","Purpose"], [
    ["NGFW (Next-Generation Firewall)","Perimeter (internet edge)",   "DPI, IPS/IDS, VPN termination \u2014 Palo Alto / Fortinet"],
    ["External Firewall",              "Between internet and DMZ",    "Filters all inbound/outbound internet traffic"],
    ["Internal Firewall",              "Between DMZ and internal LAN","Prevents lateral movement in case of breach"],
    ["Software Firewalls",             "On each client/server",       "Host-based protection \u2014 Windows Defender, iptables"],
    ["Web Application Firewall (WAF)", "In front of web/app servers", "Protects against SQL injection, XSS, other web attacks"],
])
add_body(doc,
    "A dedicated SIEM server (Splunk, IBM QRadar, or Microsoft Sentinel) handles security event aggregation "
    "and real-time threat detection. Security software: CrowdStrike Falcon (EDR), Bitdefender GravityZone "
    "(antivirus), Microsoft Defender for Office 365 (email security), Nessus (vulnerability scanning), "
    "Okta (IAM + MFA). Multi-Factor Authentication (MFA) is mandatory for all accounts."
)

add_heading(doc, "8. Enterprise and Operational Software", 2)
add_table(doc, ["Type","Products","Purpose"], [
    ["ERP",  "SAP S/4HANA, Microsoft Dynamics 365","Integrates finance, HR, procurement, and operations"],
    ["CRM",  "Salesforce, Microsoft Dynamics CRM", "Customer interactions, sales pipelines, support tickets"],
    ["HRM",  "Workday, SAP SuccessFactors",        "Employee management, payroll, recruitment, performance"],
    ["PM",   "Jira, Microsoft Project, Asana",     "Task tracking, sprint management, project planning"],
    ["ITSM", "ServiceNow, Freshservice",            "IT help desk, incident management (ITIL framework)"],
    ["BI",   "Microsoft Power BI, Tableau",        "Reporting, analytics, decision support dashboards"],
    ["SCM",  "SAP SCM, Oracle SCM",                "Vendor and procurement management"],
])

add_heading(doc, "9. Application Software", 2)
add_table(doc, ["Category","Applications"], [
    ["Office Productivity", "Microsoft 365 (Word, Excel, PowerPoint, Outlook, Teams) / Google Workspace"],
    ["Development Tools",   "Visual Studio, VS Code, IntelliJ IDEA, PyCharm, Android Studio"],
    ["Version Control",     "Git, GitHub Desktop, GitLab, Bitbucket"],
    ["DevOps / CI-CD",      "Jenkins, GitHub Actions, Docker, Kubernetes, Ansible"],
    ["Database Management", "DBeaver, MySQL Workbench, pgAdmin, MongoDB Compass"],
    ["Communication",       "Microsoft Teams, Slack, Zoom"],
    ["Remote Access / VPN", "Cisco AnyConnect, Microsoft Remote Desktop, TeamViewer"],
    ["Security Tools",      "Wireshark, Nmap, Burp Suite (for security team)"],
    ["Browsers",            "Google Chrome, Mozilla Firefox, Microsoft Edge"],
    ["Backup",              "Veeam Backup, Acronis Cyber Backup"],
])

add_heading(doc, "10. Network Administrators and Installation Timeline", 2)
add_body(doc, "Required IT Staff:", bold=True)
add_table(doc, ["Role","Count"], [
    ["Network Administrators (LAN/WAN, switches, routers)", "4\u20136"],
    ["System Administrators (servers, VMs, OS)",            "4\u20136"],
    ["Security Administrators (firewall, SIEM, endpoint)",  "2\u20133"],
    ["Database Administrators (DBAs)",                      "2\u20133"],
    ["Help Desk / IT Support Technicians (Level 1/2)",      "10\u201315"],
    ["DevOps / Cloud Engineers",                            "3\u20135"],
    ["IT Manager / CIO",                                    "1\u20132"],
    ["Total IT Staff",                                      "~26\u201340 people"],
])
add_body(doc, "Estimated Installation Timeline:", bold=True)
add_table(doc, ["Phase","Duration","Activities"], [
    ["Phase 1: Planning & Procurement", "4\u20136 weeks",  "Requirements, vendor selection, hardware/software ordering"],
    ["Phase 2: Physical Infrastructure","6\u20138 weeks",  "Cabling, rack installation, switch/router and server room setup"],
    ["Phase 3: Server & OS Setup",      "4\u20136 weeks",  "OS installation, AD/LDAP, virtualization, DNS/DHCP configuration"],
    ["Phase 4: Network Configuration",  "3\u20134 weeks",  "VLANs, routing, firewall rules, VPN, wireless deployment"],
    ["Phase 5: Software Deployment",    "4\u20136 weeks",  "ERP/CRM, productivity apps, security software, client images"],
    ["Phase 6: Testing & UAT",          "2\u20134 weeks",  "Security testing, load testing, user acceptance testing"],
    ["Phase 7: Training & Go-Live",     "2\u20133 weeks",  "Staff training, phased go-live, monitoring and issue resolution"],
    ["Total Estimated Time",            "25\u201337 weeks","Approximately 6\u20139 months for a full enterprise-grade installation"],
])
add_body(doc,
    "With a larger experienced team and parallel workstreams, the timeline could be compressed to 4\u20136 months. "
    "However, for a quality, enterprise-grade installation with minimal disruption, 6\u20139 months is a realistic estimate."
)

# Save
out = r"c:\Users\mete\AndroidStudioProjects\django\acm474_homework\ACM474_Term_Homework_KivancMeteTaskiran.docx"
doc.save(out)
print("Saved:", out)
